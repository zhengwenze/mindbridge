from __future__ import annotations

import hashlib
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterator

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader


PARSER_VERSION = "1"
PDF_PAGE_BOUNDARY_TEMPLATE = "--- PDF PAGE {page_number} ---"
SUPPORTED_DOCUMENT_EXTENSIONS = {".txt", ".md", ".markdown", ".pdf", ".docx"}

_CONTROL_CHARACTERS = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]")
_ZERO_WIDTH_CHARACTERS = re.compile(r"[\u180e\u200b-\u200d\u2060\ufeff]")
_EXCESSIVE_BLANK_LINES = re.compile(r"\n[ \t]*(?:\n[ \t]*){2,}")


class DocumentParseError(ValueError):
    """Raised when a supported document cannot produce safe, usable text."""


class UnsupportedDocumentFormatError(DocumentParseError):
    pass


@dataclass(frozen=True)
class ParsedDocument:
    text: str
    parser_name: str
    parser_version: str
    metadata: dict[str, object]
    page_count: int | None
    warnings: list[str]


def content_hash(text: str) -> str:
    """Return the stable SHA-256 of the effective UTF-8 parsed text."""

    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def clean_text(text: str) -> tuple[str, list[str]]:
    """Remove only characters that are unambiguously meaningless.

    Structural whitespace, Markdown syntax, tables, lists, and punctuation are
    intentionally left intact. Warnings are stable and emitted once per kind.
    """

    warnings: list[str] = []
    cleaned = (text or "").replace("\r\n", "\n").replace("\r", "\n")

    cleaned, control_count = _CONTROL_CHARACTERS.subn("", cleaned)
    if control_count:
        warnings.append("已移除非法控制字符")

    cleaned, zero_width_count = _ZERO_WIDTH_CHARACTERS.subn("", cleaned)
    if zero_width_count:
        warnings.append("已移除零宽字符")

    replacement_count = cleaned.count("\ufffd")
    if replacement_count:
        cleaned = cleaned.replace("\ufffd", "")
        warnings.append("已移除异常替换字符 U+FFFD")

    cleaned = _EXCESSIVE_BLANK_LINES.sub("\n\n", cleaned)
    return cleaned.strip(), warnings


def parse_document(path: Path, name: str, settings: Any) -> ParsedDocument:
    """Parse a supported local document without changing external state."""

    extension = Path(name).suffix.lower()
    if extension not in SUPPORTED_DOCUMENT_EXTENSIONS:
        raise UnsupportedDocumentFormatError("仅支持 txt、md、markdown、pdf、docx 格式")

    try:
        if extension in {".txt", ".md", ".markdown"}:
            return _parse_text_document(path, extension)
        if extension == ".pdf":
            return _parse_pdf(path)
        return _parse_docx(path, settings)
    except DocumentParseError:
        raise
    except UnicodeDecodeError as exc:
        raise DocumentParseError("TXT 和 Markdown 文档必须使用 UTF-8 或 UTF-8-SIG 编码") from exc
    except Exception as exc:
        raise DocumentParseError("文档损坏或无法解析") from exc


def _parse_text_document(path: Path, extension: str) -> ParsedDocument:
    raw_text = path.read_text(encoding="utf-8-sig")
    text, warnings = clean_text(raw_text)
    parser_name = "markdown" if extension in {".md", ".markdown"} else "text"
    return ParsedDocument(
        text=text,
        parser_name=parser_name,
        parser_version=PARSER_VERSION,
        metadata={"file_extension": extension},
        page_count=None,
        warnings=warnings,
    )


def _parse_pdf(path: Path) -> ParsedDocument:
    reader = PdfReader(str(path))
    if reader.is_encrypted:
        raise DocumentParseError("暂不支持加密 PDF")

    page_count = len(reader.pages)
    pages: list[str] = []
    empty_pages: list[int] = []
    warnings: list[str] = []
    for page_number, page in enumerate(reader.pages, start=1):
        page_text, page_warnings = clean_text(page.extract_text() or "")
        _extend_unique(warnings, page_warnings)
        if not page_text:
            empty_pages.append(page_number)
            warnings.append(f"PDF 第 {page_number} 页未提取到文本")
        pages.append(page_text)

    if not any(pages):
        raise DocumentParseError("PDF 未提取到文本，可能是扫描型 PDF；当前阶段不支持 OCR")

    text = pages[0]
    for page_number, page_text in enumerate(pages[1:], start=2):
        boundary = PDF_PAGE_BOUNDARY_TEMPLATE.format(page_number=page_number)
        text = f"{text}\n\n{boundary}\n\n{page_text}"
    text, final_warnings = clean_text(text)
    _extend_unique(warnings, final_warnings)

    return ParsedDocument(
        text=text,
        parser_name="pypdf",
        parser_version=PARSER_VERSION,
        metadata={"page_count": page_count, "empty_pages": empty_pages},
        page_count=page_count,
        warnings=warnings,
    )


def _parse_docx(path: Path, settings: Any) -> ParsedDocument:
    _validate_docx_archive(path, settings)
    document = DocxDocument(str(path))
    rendered: list[tuple[str, str]] = []
    paragraph_count = 0
    table_count = 0

    for block in _iter_docx_blocks(document):
        if isinstance(block, Paragraph):
            paragraph_count += 1
            kind, text = _render_paragraph(block, document)
            rendered.append((kind, text))
        else:
            table_count += 1
            rendered.append(("table", _render_table(block)))

    raw_text = _join_docx_blocks(rendered)
    text, warnings = clean_text(raw_text)
    return ParsedDocument(
        text=text,
        parser_name="python-docx",
        parser_version=PARSER_VERSION,
        metadata={
            "paragraph_count": paragraph_count,
            "table_count": table_count,
        },
        page_count=None,
        warnings=warnings,
    )


def _validate_docx_archive(path: Path, settings: Any) -> None:
    max_uncompressed_bytes = int(
        getattr(settings, "knowledge_docx_max_uncompressed_bytes", 200 * 1024 * 1024)
    )
    try:
        with zipfile.ZipFile(path) as archive:
            entries = archive.infolist()
            uncompressed_size = sum(entry.file_size for entry in entries)
            if len(entries) > 10_000 or uncompressed_size > max_uncompressed_bytes:
                raise DocumentParseError("DOCX 解压后的内容超过安全限制")
            if "[Content_Types].xml" not in archive.namelist():
                raise DocumentParseError("DOCX 文件结构无效")
    except zipfile.BadZipFile as exc:
        raise DocumentParseError("DOCX 文件结构无效") from exc


def _iter_docx_blocks(document: Any) -> Iterator[Paragraph | Table]:
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _render_paragraph(paragraph: Paragraph, document: Any) -> tuple[str, str]:
    text = paragraph.text.strip()
    if not text:
        return "empty", ""

    heading_level = _heading_level(paragraph)
    if heading_level is not None:
        return "heading", f"{'#' * heading_level} {text}"

    list_info = _list_info(paragraph, document)
    if list_info is not None:
        ordered, level = list_info
        marker = "1. " if ordered else "- "
        return "list", f"{'  ' * level}{marker}{text}"

    return "paragraph", text


def _heading_level(paragraph: Paragraph) -> int | None:
    style = paragraph.style
    if style is None:
        return None
    for candidate in (style.name or "", style.style_id or ""):
        match = re.fullmatch(r"Heading\s*([1-9])", candidate, flags=re.IGNORECASE)
        if match:
            return int(match.group(1))
    return None


def _list_info(paragraph: Paragraph, document: Any) -> tuple[bool, int] | None:
    style_name = (paragraph.style.name if paragraph.style is not None else "") or ""
    lowered_style = style_name.lower()

    num_properties = None
    if paragraph._p.pPr is not None:
        num_properties = paragraph._p.pPr.numPr
    if num_properties is None and paragraph.style is not None:
        style_ppr = paragraph.style.element.pPr
        num_properties = style_ppr.numPr if style_ppr is not None else None

    level = 0
    number_format: str | None = None
    if num_properties is not None:
        if num_properties.ilvl is not None and num_properties.ilvl.val is not None:
            level = max(0, int(num_properties.ilvl.val))
        if num_properties.numId is not None and num_properties.numId.val is not None:
            number_format = _numbering_format(document, int(num_properties.numId.val), level)

    if "bullet" in lowered_style:
        return False, level
    if "number" in lowered_style:
        return True, level
    if number_format is not None:
        return number_format != "bullet", level
    return None


def _numbering_format(document: Any, num_id: int, level: int) -> str | None:
    numbering = document.part.numbering_part.element
    num = next(
        (
            item
            for item in numbering.findall(qn("w:num"))
            if item.get(qn("w:numId")) == str(num_id)
        ),
        None,
    )
    if num is None:
        return None
    abstract_id_element = num.find(qn("w:abstractNumId"))
    if abstract_id_element is None:
        return None
    abstract_id = abstract_id_element.get(qn("w:val"))
    abstract_num = next(
        (
            item
            for item in numbering.findall(qn("w:abstractNum"))
            if item.get(qn("w:abstractNumId")) == abstract_id
        ),
        None,
    )
    if abstract_num is None:
        return None
    levels = abstract_num.findall(qn("w:lvl"))
    selected = next(
        (item for item in levels if item.get(qn("w:ilvl")) == str(level)),
        levels[0] if levels else None,
    )
    if selected is None:
        return None
    number_format = selected.find(qn("w:numFmt"))
    return number_format.get(qn("w:val")) if number_format is not None else None


def _render_table(table: Table) -> str:
    rows = [[_render_table_cell(cell.text) for cell in row.cells] for row in table.rows]
    column_count = len(table.columns)
    if not rows:
        return f"表格（0 行，{column_count} 列）"

    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    lines = [
        f"| {' | '.join(header)} |",
        f"| {' | '.join('---' for _ in range(width))} |",
    ]
    lines.extend(f"| {' | '.join(row)} |" for row in normalized[1:])
    return "\n".join(lines)


def _render_table_cell(value: str) -> str:
    normalized = value.replace("\r\n", "\n").replace("\r", "\n").strip()
    return normalized.replace("|", r"\|").replace("\n", "<br>")


def _join_docx_blocks(blocks: list[tuple[str, str]]) -> str:
    output = ""
    previous_kind: str | None = None
    for kind, text in blocks:
        if not output:
            output = text
        elif kind == "list" and previous_kind == "list":
            output = f"{output}\n{text}"
        else:
            output = f"{output}\n\n{text}"
        previous_kind = kind
    return output


def _extend_unique(target: list[str], values: list[str]) -> None:
    for value in values:
        if value not in target:
            target.append(value)
