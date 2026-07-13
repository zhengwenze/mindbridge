from __future__ import annotations

import json
import logging
import math
import re
import shutil
import unicodedata
import uuid
import zipfile
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path, PurePosixPath
from typing import Any, Hashable, Iterable, Iterator

from docx import Document as DocxDocument
from pypdf import PdfReader
from sqlalchemy import func
from sqlalchemy.exc import IntegrityError
from sqlalchemy.orm import Session, joinedload

from app.core.config import Settings
from app.models.entities import KnowledgeBase, KnowledgeBaseOperationLog, KnowledgeBaseReference, KnowledgeChunk, KnowledgeDocument, UserAccount
from app.services.vector_store import FALLBACK_RETRIEVAL_LABEL, PRIMARY_RETRIEVAL_LABEL, ChromaVectorStore, VectorStoreUnavailable


logger = logging.getLogger(__name__)

BASE_KNOWLEDGE_NAME = "心理健康基础知识库"
POLICY_KNOWLEDGE_NAME = "校园心理咨询政策库"
CRISIS_KNOWLEDGE_NAME = "危机干预知识库"
DEFAULT_KNOWLEDGE_BASES = (
    (BASE_KNOWLEDGE_NAME, "存储大学生常见心理健康问题、情绪调节、压力管理、人际关系与心理健康科普资料。"),
    (POLICY_KNOWLEDGE_NAME, "存储学校心理咨询预约流程、服务制度、咨询须知、保密原则与校内心理服务信息。"),
    (CRISIS_KNOWLEDGE_NAME, "存储心理危机识别、风险分级、危机干预流程、紧急转介与人工介入规范。"),
)
MANAGEABLE_STATUSES = {"active", "disabled"}
DELETING = "DELETING"
DELETE_FAILED = "DELETE_FAILED"
REFERENCE_TYPES = {"agent", "application", "department", "running_task"}
SUPPORTED_DOCUMENT_EXTENSIONS = {".txt", ".md", ".markdown", ".pdf", ".docx"}


class KnowledgeBaseError(RuntimeError):
    status_code = 400

    def __init__(self, message: str, *, detail: dict | None = None):
        super().__init__(message)
        self.detail = detail


class KnowledgeBaseNotFound(KnowledgeBaseError):
    status_code = 404


class KnowledgeBaseConflict(KnowledgeBaseError):
    status_code = 409


class KnowledgeBaseUnavailable(KnowledgeBaseError):
    status_code = 409


class KnowledgeDocumentTooLarge(KnowledgeBaseError):
    status_code = 413


class KnowledgeDocumentUnsupported(KnowledgeBaseError):
    status_code = 415


class KnowledgeDocumentInvalid(KnowledgeBaseError):
    status_code = 422


class KnowledgeDocumentProcessingError(KnowledgeBaseError):
    status_code = 503


@dataclass
class SearchResult:
    chunk_id: int | None
    source: str
    content: str
    score: float
    knowledge_base_id: int | None = None
    knowledge_base_name: str = ""
    document_id: int | None = None
    document_name: str = ""


@dataclass
class RetrievalCandidate:
    result: SearchResult
    vector_score: float = 0.0
    bm25_score: float = 0.0


class KnowledgeBaseService:
    def __init__(self, db: Session, settings: Settings):
        self.db = db
        self.settings = settings
        self.vector_store = ChromaVectorStore(settings)

    def ensure_defaults(self, actor: UserAccount | None = None) -> list[KnowledgeBase]:
        bases = []
        for name, description in DEFAULT_KNOWLEDGE_BASES:
            base = self.db.query(KnowledgeBase).filter(KnowledgeBase.name == name, KnowledgeBase.deleted_at.is_(None)).first()
            if base is None:
                base = self.create(name, description, actor, audit=False)
            elif base.created_by is None and actor is not None:
                base.created_by = actor.id
                self.db.commit()
            if base.deleted_at is None and self.vector_store.can_embed and not self.vector_store.collection_exists(base.collection_name):
                try:
                    self.vector_store.create_collection(base.collection_name)
                    if base.status == "error":
                        base.status = "active"
                    self.db.commit()
                except Exception:
                    base.status = "error"
                    self.db.commit()
            bases.append(base)
        return bases

    def create(self, name: str, description: str | None, actor: UserAccount | None, audit: bool = True) -> KnowledgeBase:
        normalized = normalize_name(name)
        if self.db.query(KnowledgeBase.id).filter(KnowledgeBase.name == normalized, KnowledgeBase.deleted_at.is_(None)).first():
            raise KnowledgeBaseConflict("知识库名称已存在")
        base = KnowledgeBase(name=normalized, description=(description or "").strip(), collection_name="pending", status="indexing", created_by=actor.id if actor else None)
        self.db.add(base)
        try:
            self.db.flush()
            base.collection_name = self.collection_name(base.id)
            self.db.commit()
        except IntegrityError as exc:
            self.db.rollback()
            raise KnowledgeBaseConflict("知识库名称已存在") from exc
        try:
            if self.vector_store.can_embed:
                self.vector_store.create_collection(base.collection_name)
            base.status = "active"
            self._audit(base.id, actor, "create", "success", {"collectionName": base.collection_name} if audit else {})
            self.db.commit()
        except Exception as exc:
            base.status = "error"
            self._audit(base.id, actor, "create", "error", {"error": safe_error(exc)} if audit else {})
            self.db.commit()
            logger.exception("knowledge base collection creation failed id=%s", base.id)
        self.db.refresh(base)
        return base

    def list(self, *, name: str | None = None, status: str | None = None, created_from: datetime | None = None, created_to: datetime | None = None, include_deleted: bool = False, page: int = 1, page_size: int = 20) -> dict:
        document_counts = self.db.query(KnowledgeDocument.knowledge_base_id.label("kb_id"), func.count(KnowledgeDocument.id).label("count")).filter(KnowledgeDocument.deleted_at.is_(None)).group_by(KnowledgeDocument.knowledge_base_id).subquery()
        chunk_counts = self.db.query(KnowledgeChunk.knowledge_base_id.label("kb_id"), func.count(KnowledgeChunk.id).label("count")).group_by(KnowledgeChunk.knowledge_base_id).subquery()
        query = self.db.query(KnowledgeBase, func.coalesce(document_counts.c.count, 0), func.coalesce(chunk_counts.c.count, 0)).outerjoin(document_counts, document_counts.c.kb_id == KnowledgeBase.id).outerjoin(chunk_counts, chunk_counts.c.kb_id == KnowledgeBase.id)
        if not include_deleted:
            query = query.filter(KnowledgeBase.deleted_at.is_(None))
        if name:
            query = query.filter(KnowledgeBase.name.ilike(f"%{name.strip()}%"))
        if status:
            query = query.filter(KnowledgeBase.status == status)
        if created_from:
            query = query.filter(KnowledgeBase.created_at >= created_from)
        if created_to:
            query = query.filter(KnowledgeBase.created_at <= created_to)
        total = query.count()
        rows = query.order_by(KnowledgeBase.created_at.desc()).offset((page - 1) * page_size).limit(page_size).all()
        return {"items": [self.serialize(base, int(documents), int(chunks)) for base, documents, chunks in rows], "total": total, "page": page, "pageSize": page_size}

    def get(self, knowledge_base_id: int, include_deleted: bool = False) -> KnowledgeBase:
        query = self.db.query(KnowledgeBase).filter(KnowledgeBase.id == knowledge_base_id)
        if not include_deleted:
            query = query.filter(KnowledgeBase.deleted_at.is_(None))
        base = query.first()
        if base is None:
            raise KnowledgeBaseNotFound("知识库不存在")
        return base

    def detail(self, knowledge_base_id: int, include_deleted: bool = False) -> dict:
        base = self.get(knowledge_base_id, include_deleted)
        document_count = self.db.query(KnowledgeDocument).filter(KnowledgeDocument.knowledge_base_id == base.id, KnowledgeDocument.deleted_at.is_(None)).count()
        chunk_count = self.db.query(KnowledgeChunk).filter(KnowledgeChunk.knowledge_base_id == base.id).count()
        return self.serialize(base, document_count, chunk_count, include_vector=True)

    def update(self, knowledge_base_id: int, *, name: str | None, description: str | None, status: str | None, actor: UserAccount) -> KnowledgeBase:
        base = self.get(knowledge_base_id)
        if base.status not in MANAGEABLE_STATUSES:
            raise KnowledgeBaseUnavailable("知识库正在处理或不可用，不能编辑")
        if status is not None and status not in MANAGEABLE_STATUSES:
            raise KnowledgeBaseError("状态仅允许设为 active 或 disabled")
        if name is not None:
            normalized = normalize_name(name)
            conflict = self.db.query(KnowledgeBase.id).filter(KnowledgeBase.name == normalized, KnowledgeBase.deleted_at.is_(None), KnowledgeBase.id != base.id).first()
            if conflict:
                raise KnowledgeBaseConflict("知识库名称已存在")
            base.name = normalized
        if description is not None:
            base.description = description.strip()
        if status is not None:
            base.status = status
        try:
            self._audit(base.id, actor, "update", "success", {"collectionName": base.collection_name})
            self.db.commit()
        except IntegrityError as exc:
            self.db.rollback()
            raise KnowledgeBaseConflict("知识库名称已存在") from exc
        self.db.refresh(base)
        return base

    def delete(self, knowledge_base_id: int, actor: UserAccount) -> dict:
        base = self.db.query(KnowledgeBase).filter(KnowledgeBase.id == knowledge_base_id).with_for_update().first()
        if base is None:
            completed = self.db.query(KnowledgeBaseOperationLog.id).filter(
                KnowledgeBaseOperationLog.knowledge_base_id == knowledge_base_id,
                KnowledgeBaseOperationLog.action == "delete",
                KnowledgeBaseOperationLog.status == "success",
            ).first()
            if completed:
                return {"id": knowledge_base_id, "status": "DELETED", "idempotent": True}
            raise KnowledgeBaseNotFound("知识库不存在")

        references = self.blocking_references(base.id)
        if references:
            detail = {"code": "KNOWLEDGE_BASE_IN_USE", "message": "知识库仍被引用，无法删除", "references": references}
            self._audit(base.id, actor, "delete", "blocked", detail)
            self.db.commit()
            raise KnowledgeBaseConflict(detail["message"], detail=detail)

        if base.status == DELETING:
            return {"id": base.id, "status": DELETING, "idempotent": True}
        if not self._isolated_collection(base):
            detail = {"code": "COLLECTION_NOT_ISOLATED", "message": "知识库尚未完成 Chroma collection 隔离，禁止删除"}
            self._audit(base.id, actor, "delete", "blocked", detail)
            self.db.commit()
            raise KnowledgeBaseConflict(detail["message"], detail=detail)

        # This committed compare-and-set is the deletion barrier: all writers
        # and retrieval paths only accept active bases after this point.
        updated = self.db.query(KnowledgeBase).filter(
            KnowledgeBase.id == base.id,
            KnowledgeBase.status.in_(["active", "disabled", "error", DELETE_FAILED]),
        ).update({KnowledgeBase.status: DELETING}, synchronize_session=False)
        if updated != 1:
            self.db.rollback()
            raise KnowledgeBaseConflict("知识库当前状态不允许删除")
        self._audit(base.id, actor, "delete", "started", {"retry": base.status == DELETE_FAILED})
        self.db.commit()
        base = self.db.get(KnowledgeBase, knowledge_base_id)
        try:
            if not self.vector_store.can_embed:
                raise VectorStoreUnavailable("Chroma 不可用，无法确认 collection 已安全删除")
            self.vector_store.delete_collection(base.collection_name)
            remove_knowledge_base_files(base.id, self.settings)
            self.db.query(KnowledgeBaseReference).filter(KnowledgeBaseReference.knowledge_base_id == base.id).delete()
            self.db.delete(base)
            self._audit(knowledge_base_id, actor, "delete", "success", {"collectionName": base.collection_name})
            self.db.commit()
            return {"id": knowledge_base_id, "status": "DELETED", "idempotent": False}
        except Exception as exc:
            self.db.rollback()
            base = self.db.get(KnowledgeBase, knowledge_base_id)
            if base is not None:
                base.status = DELETE_FAILED
            self._audit(knowledge_base_id, actor, "delete", "failed", {"error": safe_error(exc)})
            self.db.commit()
            raise KnowledgeBaseError("知识库删除失败，已进入 DELETE_FAILED，可再次点击删除重试") from exc

    def register_reference(
        self,
        knowledge_base_id: int,
        *,
        reference_type: str,
        reference_id: str,
        reference_name: str = "",
        status: str = "active",
        blocking: bool = True,
    ) -> KnowledgeBaseReference:
        if reference_type not in REFERENCE_TYPES:
            raise KnowledgeBaseError("不支持的知识库引用类型")
        base = self.db.query(KnowledgeBase).filter(KnowledgeBase.id == knowledge_base_id).with_for_update().first()
        if base is None:
            raise KnowledgeBaseNotFound("知识库不存在")
        if base.status not in MANAGEABLE_STATUSES:
            raise KnowledgeBaseUnavailable("知识库正在删除或不可用，不能新增引用")
        row = self.db.query(KnowledgeBaseReference).filter_by(
            knowledge_base_id=base.id, reference_type=reference_type, reference_id=str(reference_id)
        ).first()
        if row is None:
            row = KnowledgeBaseReference(
                knowledge_base_id=base.id,
                reference_type=reference_type,
                reference_id=str(reference_id),
            )
            self.db.add(row)
        row.reference_name = reference_name.strip()
        row.status = status
        row.blocking = blocking
        self.db.commit()
        self.db.refresh(row)
        return row

    def release_reference(self, knowledge_base_id: int, *, reference_type: str, reference_id: str) -> None:
        self.db.query(KnowledgeBaseReference).filter_by(
            knowledge_base_id=knowledge_base_id, reference_type=reference_type, reference_id=str(reference_id)
        ).delete()
        self.db.commit()

    def status(self, knowledge_base_id: int) -> dict:
        base = self.get(knowledge_base_id)
        exists = self.vector_store.can_embed and self.vector_store.collection_exists(base.collection_name)
        vector_count = self.vector_store.count(base.collection_name) if exists else 0
        return {**self.detail(base.id), "collectionExists": exists, "vectorCount": vector_count, "embeddingModel": self.settings.embedding_model, "vectorError": self.vector_store.error}

    def default_retrieval_ids(self, query: str) -> list[int]:
        active = self.db.query(KnowledgeBase).filter(KnowledgeBase.deleted_at.is_(None), KnowledgeBase.status == "active").all()
        mapping = {base.name: base.id for base in active}
        ids = [mapping[BASE_KNOWLEDGE_NAME]] if BASE_KNOWLEDGE_NAME in mapping else []
        text = query.lower()
        if any(term in text for term in ("预约", "咨询流程", "保密", "制度", "政策", "校内服务")) and POLICY_KNOWLEDGE_NAME in mapping:
            ids.append(mapping[POLICY_KNOWLEDGE_NAME])
        if any(term in text for term in ("自杀", "自伤", "危机", "紧急", "转介", "生命")) and CRISIS_KNOWLEDGE_NAME in mapping:
            ids.append(mapping[CRISIS_KNOWLEDGE_NAME])
        return ids

    def collection_name(self, knowledge_base_id: int) -> str:
        return f"mindbridge_kb_{knowledge_base_id}"

    def document_count(self, base_id: int) -> int:
        return self.db.query(KnowledgeDocument).filter(KnowledgeDocument.knowledge_base_id == base_id, KnowledgeDocument.deleted_at.is_(None)).count()

    def chunk_count(self, base_id: int) -> int:
        return self.db.query(KnowledgeChunk).filter(KnowledgeChunk.knowledge_base_id == base_id).count()

    def blocking_references(self, base_id: int) -> list[dict]:
        rows = self.db.query(KnowledgeBaseReference).filter(
            KnowledgeBaseReference.knowledge_base_id == base_id,
            KnowledgeBaseReference.blocking.is_(True),
            KnowledgeBaseReference.status.in_(["active", "running"]),
            KnowledgeBaseReference.reference_type.in_(REFERENCE_TYPES),
        ).order_by(KnowledgeBaseReference.reference_type, KnowledgeBaseReference.id).all()
        return [{"type": row.reference_type, "id": row.reference_id, "name": row.reference_name, "status": row.status} for row in rows]

    def _isolated_collection(self, base: KnowledgeBase) -> bool:
        return base.collection_name == self.collection_name(base.id) and base.collection_name != "mindbridge_knowledge"

    def serialize(self, base: KnowledgeBase, document_count: int, chunk_count: int, include_vector: bool = False) -> dict:
        result = {"id": base.id, "name": base.name, "description": base.description, "collectionName": base.collection_name, "status": base.status, "createdBy": base.created_by, "createdAt": base.created_at, "updatedAt": base.updated_at, "deletedAt": base.deleted_at, "documentCount": document_count, "chunkCount": chunk_count}
        if include_vector:
            exists = self.vector_store.can_embed and self.vector_store.collection_exists(base.collection_name)
            result.update({"collectionExists": exists, "vectorCount": self.vector_store.count(base.collection_name) if exists else 0})
        return result

    def _audit(self, base_id: int | None, actor: UserAccount | None, action: str, status: str, detail: dict) -> None:
        self.db.add(KnowledgeBaseOperationLog(knowledge_base_id=base_id, actor_id=actor.id if actor else None, action=action, status=status, detail_json=json.dumps(detail, ensure_ascii=False)))


class KnowledgeDocumentService:
    def __init__(self, db: Session, settings: Settings):
        self.db = db
        self.settings = settings
        self.bases = KnowledgeBaseService(db, settings)
        self.vector_store = self.bases.vector_store

    def ingest_file(
        self,
        knowledge_base_id: int,
        filename: str,
        data: bytes,
        actor: UserAccount | None = None,
        *,
        relative_path: str | None = None,
        replace_existing: bool = False,
    ) -> dict:
        if len(data) > self.settings.knowledge_upload_max_bytes:
            raise KnowledgeDocumentTooLarge("单个文件不能超过 50 MB")
        temp_path = create_upload_temp(self.settings, Path(filename).suffix)
        try:
            temp_path.write_bytes(data)
            return self.ingest_path(
                knowledge_base_id,
                filename,
                relative_path,
                temp_path,
                len(data),
                actor,
                replace_existing=replace_existing,
            )
        finally:
            temp_path.unlink(missing_ok=True)

    def ingest_path(
        self,
        knowledge_base_id: int,
        filename: str,
        relative_path: str | None,
        temp_path: Path,
        file_size: int,
        actor: UserAccount | None = None,
        *,
        replace_existing: bool = False,
    ) -> dict:
        base = self.bases.get(knowledge_base_id)
        self._assert_writable(base)
        normalized_path = normalize_relative_path(relative_path, filename)
        safe_name = PurePosixPath(normalized_path).name
        validate_document_extension(safe_name)
        existing = self.db.query(KnowledgeDocument).filter(
            KnowledgeDocument.knowledge_base_id == base.id,
            KnowledgeDocument.relative_path == normalized_path,
        ).first()
        if existing is not None and not replace_existing:
            raise KnowledgeBaseConflict("该知识库中已存在相同路径的文档")

        chunks = list(extract_document_chunks(temp_path, safe_name, self.settings))
        if not chunks:
            raise KnowledgeDocumentInvalid("文档没有可索引文本")

        if existing is not None:
            if self.vector_store.can_embed:
                self.vector_store.delete_document(base.collection_name, existing.id)
            remove_stored_file(existing.storage_path, self.settings)
            self.db.delete(existing)
            self.db.flush()

        document = KnowledgeDocument(
            knowledge_base_id=base.id,
            file_name=safe_name,
            relative_path=normalized_path,
            file_type=file_type(safe_name),
            file_size=file_size,
            index_status="indexing",
        )
        stored_path: Path | None = None
        self.db.add(document)
        try:
            self.db.flush()
            rows = [
                KnowledgeChunk(
                    knowledge_base_id=base.id,
                    document_id=document.id,
                    document=document,
                    source=normalized_path,
                    source_index=index,
                    content=chunk,
                )
                for index, chunk in enumerate(chunks)
            ]
            self.db.add_all(rows)
            self.db.flush()
            self._index(base, rows)
            stored_path = move_upload_to_storage(temp_path, base.id, safe_name, self.settings)
            document.storage_path = str(stored_path.relative_to(self.settings.project_root))
            document.index_status = "active"
            self.bases._audit(
                base.id,
                actor,
                "upload",
                "success",
                {"documentId": document.id, "relativePath": normalized_path, "chunks": len(rows)},
            )
            self.db.commit()
        except IntegrityError as exc:
            self._rollback_upload(base, document, stored_path)
            raise KnowledgeBaseConflict("该知识库中已存在相同路径的文档") from exc
        except Exception as exc:
            self._rollback_upload(base, document, stored_path)
            self._audit_failed_upload(base.id, actor, normalized_path, exc)
            raise KnowledgeDocumentProcessingError("文档处理失败，请检查解析或向量服务后重试") from exc
        return {
            "id": document.id,
            "knowledgeBaseId": base.id,
            "fileName": document.file_name,
            "relativePath": document.relative_path,
            "fileSize": document.file_size,
            "chunks": len(rows),
            "indexStatus": document.index_status,
        }

    def rebuild(self, knowledge_base_id: int, actor: UserAccount | None = None) -> int:
        base = self.bases.get(knowledge_base_id)
        self._assert_writable(base)
        base.status = "indexing"
        self.db.commit()
        chunks = self.db.query(KnowledgeChunk).options(joinedload(KnowledgeChunk.document)).filter(KnowledgeChunk.knowledge_base_id == base.id).order_by(KnowledgeChunk.document_id, KnowledgeChunk.source_index).all()
        try:
            self._sync(base, chunks)
            self.db.query(KnowledgeDocument).filter(KnowledgeDocument.knowledge_base_id == base.id).update({KnowledgeDocument.index_status: "active", KnowledgeDocument.error_message: ""})
            base.status = "active"
            self.bases._audit(base.id, actor, "rebuild", "success", {"chunks": len(chunks)})
            self.db.commit()
            return len(chunks)
        except Exception as exc:
            base.status = "error"
            self.bases._audit(base.id, actor, "rebuild", "error", {"error": safe_error(exc)})
            self.db.commit()
            raise KnowledgeBaseError("知识库重建失败") from exc

    def ensure_source(self, knowledge_base_id: int, source: str, content: str) -> int:
        normalized_path = normalize_relative_path(source, source)
        existing = self.db.query(KnowledgeDocument).filter(
            KnowledgeDocument.knowledge_base_id == knowledge_base_id,
            KnowledgeDocument.relative_path == normalized_path,
        ).first()
        chunks = chunk_text(content, self.settings.knowledge_chunk_size, self.settings.knowledge_chunk_overlap)
        if existing and [row.content for row in existing.chunks] == chunks:
            return len(chunks)
        return int(self.ingest_file(knowledge_base_id, source, content.encode("utf-8"), replace_existing=True)["chunks"])

    def _index(self, base: KnowledgeBase, chunks: list[KnowledgeChunk]) -> None:
        if not self.vector_store.can_embed:
            return
        batch_size = max(1, self.settings.knowledge_embedding_batch_size)
        for start in range(0, len(chunks), batch_size):
            batch = chunks[start:start + batch_size]
            embeddings = self.vector_store.embed_texts([chunk.content for chunk in batch])
            for chunk, embedding in zip(batch, embeddings):
                chunk.embedding_json = json.dumps(embedding, separators=(",", ":"))
            self.vector_store.upsert_chunks(base.collection_name, batch, embeddings)

    def _sync(self, base: KnowledgeBase, chunks: list[KnowledgeChunk]) -> None:
        if not self.vector_store.can_embed:
            return
        self._index(base, chunks)
        self.vector_store.prune_chunks(base.collection_name, chunks)

    def _rollback_upload(self, base: KnowledgeBase, document: KnowledgeDocument, stored_path: Path | None) -> None:
        document_id = document.id
        self.db.rollback()
        if stored_path is not None:
            stored_path.unlink(missing_ok=True)
        if document_id is not None and self.vector_store.can_embed:
            try:
                self.vector_store.delete_document(base.collection_name, document_id)
            except Exception:
                logger.exception("failed to compensate vectors for document id=%s", document_id)

    def _audit_failed_upload(
        self,
        knowledge_base_id: int,
        actor: UserAccount | None,
        relative_path: str,
        exc: Exception,
    ) -> None:
        try:
            self.bases._audit(
                knowledge_base_id,
                actor,
                "upload",
                "error",
                {"relativePath": relative_path, "error": safe_error(exc)},
            )
            self.db.commit()
        except Exception:
            self.db.rollback()
            logger.exception("failed to write document upload audit")

    @staticmethod
    def _assert_writable(base: KnowledgeBase) -> None:
        if base.status != "active":
            raise KnowledgeBaseUnavailable("知识库不是 active 状态，不能上传或重建")


class KnowledgeService:
    """RAG façade. Retrieval is always constrained to trusted knowledge-base IDs."""

    def __init__(self, db: Session, settings: Settings):
        self.db = db
        self.settings = settings
        self.bases = KnowledgeBaseService(db, settings)
        self.documents = KnowledgeDocumentService(db, settings)
        self.vector_store = self.bases.vector_store

    def retrieve(self, query: str, top_k: int | None = None, knowledge_base_ids: Iterable[int] | None = None) -> list[SearchResult]:
        top_k = top_k or self.settings.knowledge_top_k
        requested = list(knowledge_base_ids) if knowledge_base_ids is not None else self.bases.default_retrieval_ids(query)
        active = self.db.query(KnowledgeBase).filter(KnowledgeBase.id.in_(requested), KnowledgeBase.deleted_at.is_(None), KnowledgeBase.status == "active").all() if requested else []
        if not active:
            return []
        ids = [base.id for base in active]
        chunks = self.db.query(KnowledgeChunk).options(joinedload(KnowledgeChunk.document)).filter(KnowledgeChunk.knowledge_base_id.in_(ids)).all()
        vector_results = self._retrieve_vector(query, active, max(top_k, self.settings.knowledge_candidate_k))
        bm25_results = self._retrieve_bm25(chunks, query, max(top_k, self.settings.knowledge_candidate_k), {base.id: base for base in active})
        ranked = self._fuse_and_rerank(query, vector_results, bm25_results, top_k)
        return self._expand_best(ranked, top_k)

    def _retrieve_vector(self, query: str, bases: list[KnowledgeBase], top_k: int) -> list[SearchResult]:
        if not self.vector_store.can_embed:
            return []
        try:
            embedding = self.vector_store.embed_texts([query])[0]
            results = []
            for base in bases:
                for hit in self.vector_store.query(base.collection_name, embedding, top_k):
                    chunk = self.db.get(KnowledgeChunk, hit.chunk_id) if hit.chunk_id else None
                    if chunk is None or chunk.knowledge_base_id != base.id:
                        continue
                    document = self.db.get(KnowledgeDocument, chunk.document_id)
                    results.append(SearchResult(chunk.id, chunk.source, chunk.content, hit.score, base.id, base.name, chunk.document_id, document.file_name if document else chunk.source))
            return sorted(results, key=lambda item: item.score, reverse=True)[:top_k]
        except Exception as exc:
            self._handle_vector_error("retrieve", exc)
            return []

    def _retrieve_bm25(self, chunks: list[KnowledgeChunk], query: str, top_k: int, bases: dict[int, KnowledgeBase]) -> list[SearchResult]:
        scores = bm25_scores(query, chunks)
        rows = []
        for chunk in chunks:
            score = scores.get(chunk.id, 0.0)
            if chunk.id is not None and score > 0:
                base = bases[chunk.knowledge_base_id]
                rows.append(SearchResult(chunk.id, chunk.source, chunk.content, score, base.id, base.name, chunk.document_id, chunk.document.file_name if chunk.document else chunk.source))
        return sorted(rows, key=lambda item: item.score, reverse=True)[:top_k]

    def _fuse_and_rerank(self, query: str, vector_results: list[SearchResult], bm25_results: list[SearchResult], top_k: int) -> list[SearchResult]:
        candidates: dict[Hashable, RetrievalCandidate] = {}
        vector_scores = normalize_scores({result_key(item): item.score for item in vector_results if item.score > 0})
        bm25_scores_by_key = normalize_scores({result_key(item): item.score for item in bm25_results if item.score > 0})
        for item in [*vector_results, *bm25_results]:
            key = result_key(item)
            candidate = candidates.setdefault(key, RetrievalCandidate(item))
            candidate.vector_score = max(candidate.vector_score, vector_scores.get(key, 0.0))
            candidate.bm25_score = max(candidate.bm25_score, bm25_scores_by_key.get(key, 0.0))
        vector_weight = max(0.0, self.settings.knowledge_hybrid_vector_weight) if vector_results else 0.0
        bm25_weight = max(0.0, self.settings.knowledge_hybrid_bm25_weight)
        total = vector_weight + bm25_weight or 1.0
        results = [replace_score(item.result, (item.vector_score * vector_weight + item.bm25_score * bm25_weight) / total) for item in candidates.values()]
        results.sort(key=lambda item: item.score, reverse=True)
        if self.settings.knowledge_rerank_enabled:
            results = [replace_score(item, rerank_score(query, item.content, item.score)) for item in results]
            results.sort(key=lambda item: item.score, reverse=True)
        return results[:top_k]

    def _expand_best(self, ranked: list[SearchResult], top_k: int) -> list[SearchResult]:
        if not ranked:
            return []
        best = self._expand(ranked[0])
        return [best, *[item for item in ranked[1:] if item.chunk_id != best.chunk_id]][:top_k]

    def _expand(self, result: SearchResult) -> SearchResult:
        if result.chunk_id is None:
            return result
        chunk = self.db.get(KnowledgeChunk, result.chunk_id)
        if chunk is None:
            return result
        neighbors = self.db.query(KnowledgeChunk).filter(KnowledgeChunk.document_id == chunk.document_id, KnowledgeChunk.source_index >= max(0, chunk.source_index - 1), KnowledgeChunk.source_index <= chunk.source_index + 1).order_by(KnowledgeChunk.source_index).all()
        return SearchResult(chunk.id, chunk.source, "\n\n".join(item.content for item in neighbors), result.score, result.knowledge_base_id, result.knowledge_base_name, chunk.document_id, result.document_name)

    def _handle_vector_error(self, action: str, exc: Exception) -> None:
        if self.settings.knowledge_vector_required:
            raise exc
        logger.warning("%s %s failed; falling back to %s: %s", PRIMARY_RETRIEVAL_LABEL, action, FALLBACK_RETRIEVAL_LABEL, exc)


def normalize_name(value: str) -> str:
    value = (value or "").strip()
    if not value or len(value) > 128:
        raise KnowledgeBaseError("知识库名称不能为空且不能超过 128 个字符")
    return value


def safe_file_name(value: str) -> str:
    name = Path(value or "uploaded-file").name.strip()
    if not name or name in {".", ".."}:
        raise KnowledgeBaseError("文件名无效")
    return re.sub(r"[^A-Za-z0-9._\-\u4e00-\u9fff]", "_", name)[:256]


def file_type(name: str) -> str:
    return name.rsplit(".", 1)[-1].lower() if "." in name else "text"


def normalize_relative_path(value: str | None, fallback: str) -> str:
    raw = unicodedata.normalize("NFC", (value or fallback or "").strip()).replace("\\", "/")
    if not raw or raw.startswith("/") or re.match(r"^[A-Za-z]:", raw):
        raise KnowledgeDocumentInvalid("文档相对路径无效")
    raw_parts = raw.split("/")
    if any(part in {"", ".", ".."} for part in raw_parts):
        raise KnowledgeDocumentInvalid("文档相对路径不能包含空目录、. 或 ..")
    if any(any(ord(character) < 32 for character in part) for part in raw_parts):
        raise KnowledgeDocumentInvalid("文档相对路径包含非法字符")
    normalized = str(PurePosixPath(*raw_parts))
    if len(normalized) > 512 or len(PurePosixPath(normalized).name) > 256:
        raise KnowledgeDocumentInvalid("文档相对路径不能超过 512 字符，文件名不能超过 256 字符")
    return normalized


def validate_document_extension(name: str) -> None:
    if Path(name).suffix.lower() not in SUPPORTED_DOCUMENT_EXTENSIONS:
        raise KnowledgeDocumentUnsupported("仅支持 txt、md、markdown、pdf、docx 格式")


def create_upload_temp(settings: Settings, suffix: str = "") -> Path:
    root = settings.project_root / "data" / "knowledge-files" / ".tmp"
    root.mkdir(parents=True, exist_ok=True)
    safe_suffix = suffix.lower() if suffix.lower() in SUPPORTED_DOCUMENT_EXTENSIONS else ""
    return root / f"{uuid.uuid4().hex}{safe_suffix}"


async def receive_upload(file: Any, settings: Settings) -> tuple[Path, int]:
    """Stream a multipart upload to a staging file with a hard byte limit."""
    temp_path = create_upload_temp(settings, Path(file.filename or "").suffix)
    total = 0
    block_size = max(1, settings.knowledge_upload_read_chunk_bytes)
    try:
        with temp_path.open("wb") as target:
            while True:
                block = await file.read(block_size)
                if not block:
                    break
                total += len(block)
                if total > settings.knowledge_upload_max_bytes:
                    raise KnowledgeDocumentTooLarge("单个文件不能超过 50 MB")
                target.write(block)
        return temp_path, total
    except Exception:
        temp_path.unlink(missing_ok=True)
        raise


def move_upload_to_storage(
    temp_path: Path,
    knowledge_base_id: int,
    name: str,
    settings: Settings,
) -> Path:
    root = settings.project_root / "data" / "knowledge-files" / str(knowledge_base_id)
    root.mkdir(parents=True, exist_ok=True)
    destination = root / f"{uuid.uuid4().hex}_{safe_file_name(name)}"
    return Path(shutil.move(str(temp_path), destination))


def remove_stored_file(value: str | None, settings: Settings) -> None:
    if not value:
        return
    root = (settings.project_root / "data" / "knowledge-files").resolve()
    path = (settings.project_root / value).resolve()
    if path.is_relative_to(root) and path.is_file():
        path.unlink()


def remove_knowledge_base_files(knowledge_base_id: int, settings: Settings) -> None:
    root = (settings.project_root / "data" / "knowledge-files").resolve()
    directory = (root / str(knowledge_base_id)).resolve()
    if directory.parent != root:
        raise KnowledgeBaseError("知识库文件目录越界")
    if directory.exists():
        shutil.rmtree(directory)


def safe_error(exc: Exception) -> str:
    return re.sub(r"[\r\n]+", " ", str(exc))[:300] or type(exc).__name__


def extract_document_chunks(path: Path, name: str, settings: Settings) -> Iterator[str]:
    try:
        extension = Path(name).suffix.lower()
        if extension in {".txt", ".md", ".markdown"}:
            blocks = iter_text_blocks(path)
        elif extension == ".pdf":
            blocks = iter_pdf_blocks(path)
        elif extension == ".docx":
            blocks = iter_docx_blocks(path, settings)
        else:
            raise KnowledgeDocumentUnsupported("不支持的文档格式")
        yield from chunk_blocks(blocks, settings.knowledge_chunk_size, settings.knowledge_chunk_overlap)
    except KnowledgeBaseError:
        raise
    except UnicodeDecodeError as exc:
        raise KnowledgeDocumentInvalid("TXT 和 Markdown 文档必须使用 UTF-8 编码") from exc
    except Exception as exc:
        raise KnowledgeDocumentInvalid("文档损坏或无法解析") from exc


def iter_text_blocks(path: Path) -> Iterator[str]:
    with path.open("r", encoding="utf-8-sig") as source:
        while block := source.read(64 * 1024):
            yield block


def iter_pdf_blocks(path: Path) -> Iterator[str]:
    reader = PdfReader(str(path))
    if reader.is_encrypted:
        raise KnowledgeDocumentInvalid("暂不支持加密 PDF")
    for page in reader.pages:
        yield page.extract_text() or ""


def iter_docx_blocks(path: Path, settings: Settings) -> Iterator[str]:
    with zipfile.ZipFile(path) as archive:
        entries = archive.infolist()
        uncompressed_size = sum(entry.file_size for entry in entries)
        if len(entries) > 10_000 or uncompressed_size > settings.knowledge_docx_max_uncompressed_bytes:
            raise KnowledgeDocumentInvalid("DOCX 解压后的内容超过安全限制")
        if "[Content_Types].xml" not in archive.namelist():
            raise KnowledgeDocumentInvalid("DOCX 文件结构无效")
    document = DocxDocument(str(path))
    for paragraph in document.paragraphs:
        yield paragraph.text
    for table in document.tables:
        for row in table.rows:
            yield "\t".join(cell.text for cell in row.cells)


def chunk_blocks(blocks: Iterable[str], size: int, overlap: int) -> Iterator[str]:
    chunk_size = max(1, size)
    chunk_overlap = min(max(0, overlap), chunk_size - 1)
    step = chunk_size - chunk_overlap
    buffer = ""
    emitted = False
    for block in blocks:
        normalized = re.sub(r"\s+", " ", block or "").strip()
        if not normalized:
            continue
        buffer = f"{buffer} {normalized}".strip() if buffer else normalized
        while len(buffer) >= chunk_size:
            yield buffer[:chunk_size]
            emitted = True
            buffer = buffer[step:]
    if buffer and (not emitted or len(buffer) > chunk_overlap):
        yield buffer


def chunk_text(content: str, size: int, overlap: int) -> list[str]:
    return list(chunk_blocks([content], size, overlap))


def bm25_scores(query: str, chunks: list[KnowledgeChunk]) -> dict[int, float]:
    query_terms = counts(tokenize(query))
    if not query_terms or not chunks:
        return {}
    documents, doc_freqs = [], {}
    for chunk in chunks:
        if chunk.id is None:
            continue
        terms = counts(tokenize(chunk.content))
        documents.append((chunk.id, terms, sum(terms.values())))
        for term in terms:
            doc_freqs[term] = doc_freqs.get(term, 0) + 1
    avg_len = sum(length for _, _, length in documents) / max(1, len(documents))
    scores = {}
    for chunk_id, terms, length in documents:
        score = 0.0
        for term, frequency in query_terms.items():
            tf = terms.get(term, 0)
            if not tf:
                continue
            idf = math.log(1 + (len(documents) - doc_freqs.get(term, 0) + 0.5) / (doc_freqs.get(term, 0) + 0.5))
            score += frequency * idf * (tf * 2.2) / (tf + 1.2 * (1 - 0.75 + 0.75 * length / max(1, avg_len)))
        scores[chunk_id] = score
    return scores


def tokenize(text: str) -> list[str]:
    words = re.findall(r"[A-Za-z0-9_]+|[\u4e00-\u9fff]", (text or "").lower())
    chinese = "".join(token for token in words if len(token) == 1 and "\u4e00" <= token <= "\u9fff")
    return words + [chinese[index:index + 2] for index in range(max(0, len(chinese) - 1))]


def counts(items: list[str]) -> dict[str, int]:
    result: dict[str, int] = {}
    for item in items:
        result[item] = result.get(item, 0) + 1
    return result


def normalize_scores(scores: dict[Hashable, float]) -> dict[Hashable, float]:
    if not scores:
        return {}
    maximum = max(scores.values())
    return {key: value / maximum if maximum else 0.0 for key, value in scores.items()}


def result_key(result: SearchResult) -> Hashable:
    return result.knowledge_base_id, result.chunk_id or (result.source, result.content)


def replace_score(result: SearchResult, score: float) -> SearchResult:
    return SearchResult(result.chunk_id, result.source, result.content, score, result.knowledge_base_id, result.knowledge_base_name, result.document_id, result.document_name)


def rerank_score(query: str, content: str, base: float) -> float:
    query_tokens = set(tokenize(query))
    content_tokens = set(tokenize(content))
    return base * 0.85 + (len(query_tokens & content_tokens) / max(1, len(query_tokens))) * 0.15
