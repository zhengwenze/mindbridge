from __future__ import annotations

import tempfile
import unittest
from pathlib import Path
from types import SimpleNamespace

from docx import Document as DocxDocument

from app.services.document_parser import (
    DocumentParseError,
    clean_text,
    content_hash,
    parse_document,
)
from app.services.document_splitter import (
    SplitterConfig,
    SplitterConfigError,
    split_text,
    validate_splitter_config,
)


def parser_settings(**overrides):
    values = {"knowledge_docx_max_uncompressed_bytes": 200 * 1024 * 1024}
    values.update(overrides)
    return SimpleNamespace(**values)


def write_text_pdf(path: Path, pages: list[list[str] | None]) -> None:
    """Write a small multi-page PDF without adding a test-only dependency."""

    font_id = 3
    kids: list[str] = []
    objects: dict[int, bytes] = {
        1: b"<< /Type /Catalog /Pages 2 0 R >>",
        font_id: b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    }
    for index, lines in enumerate(pages):
        page_id = 4 + index * 2
        content_id = page_id + 1
        kids.append(f"{page_id} 0 R")
        objects[page_id] = (
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 200] "
            f"/Resources << /Font << /F1 {font_id} 0 R >> >> /Contents {content_id} 0 R >>"
        ).encode("ascii")
        commands: list[str] = []
        for line_number, line in enumerate(lines or []):
            escaped = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
            if line_number:
                commands.append("0 -14 Td")
            commands.append(f"({escaped}) Tj")
        stream = f"BT /F1 12 Tf 72 160 Td {' '.join(commands)} ET".encode("ascii") if commands else b""
        objects[content_id] = (
            f"<< /Length {len(stream)} >>\nstream\n".encode("ascii")
            + stream
            + b"\nendstream"
        )

    objects[2] = (
        f"<< /Type /Pages /Kids [{' '.join(kids)}] /Count {len(pages)} >>"
    ).encode("ascii")
    output = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for object_id in range(1, max(objects) + 1):
        offsets.append(len(output))
        output.extend(f"{object_id} 0 obj\n".encode("ascii"))
        output.extend(objects[object_id])
        output.extend(b"\nendobj\n")
    xref_offset = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n".encode("ascii")
    )
    path.write_bytes(output)


class TextDocumentParserTests(unittest.TestCase):
    def test_utf8_sig_markdown_preserves_structure_and_chinese_punctuation(self):
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "guide.md"
            path.write_bytes(
                b"\xef\xbb\xbf"
                + "# 标题\r\n\r\n- 第一项\r\n- 第二项\r\n\r\n这是中文。真的！\r\n".encode("utf-8")
            )

            parsed = parse_document(path, path.name, parser_settings())

            self.assertEqual(parsed.parser_name, "markdown")
            self.assertEqual(parsed.page_count, None)
            self.assertEqual(parsed.warnings, [])
            self.assertEqual(
                parsed.text,
                "# 标题\n\n- 第一项\n- 第二项\n\n这是中文。真的！",
            )
            self.assertEqual(content_hash(parsed.text), content_hash(parsed.text))
            self.assertEqual(len(content_hash(parsed.text)), 64)

    def test_minimal_cleaning_removes_only_invalid_characters(self):
        text = "\x00# 标题\r\n\r\n\r\n\r\n- 列表\u200b\n中文，标点。\ufffd\t保留"

        cleaned, warnings = clean_text(text)

        self.assertEqual(cleaned, "# 标题\n\n- 列表\n中文，标点。\t保留")
        self.assertEqual(
            warnings,
            ["已移除非法控制字符", "已移除零宽字符", "已移除异常替换字符 U+FFFD"],
        )

    def test_rejects_non_utf8_text(self):
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "invalid.txt"
            path.write_bytes(b"\xff\xfe\xfa")

            with self.assertRaisesRegex(DocumentParseError, "UTF-8"):
                parse_document(path, path.name, parser_settings())


class DocxDocumentParserTests(unittest.TestCase):
    def test_preserves_body_order_and_converts_heading_lists_and_table(self):
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "structured.docx"
            document = DocxDocument()
            document.add_heading("中文标题", level=1)
            document.add_paragraph("标题后的普通段落。")
            document.add_paragraph("无序项目", style="List Bullet")
            document.add_paragraph("有序项目", style="List Number")
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "名称"
            table.cell(0, 1).text = "说明"
            table.cell(1, 0).text = "Mind|Bridge"
            table.cell(1, 1).text = "第一行\n第二行"
            document.add_paragraph("表格后的段落。")
            document.add_paragraph("")
            document.add_paragraph("")
            document.add_paragraph("")
            document.save(path)

            parsed = parse_document(path, path.name, parser_settings())

            self.assertEqual(parsed.parser_name, "python-docx")
            self.assertEqual(parsed.metadata["table_count"], 1)
            self.assertIn("# 中文标题", parsed.text)
            self.assertIn("- 无序项目\n1. 有序项目", parsed.text)
            self.assertIn("| 名称 | 说明 |", parsed.text)
            self.assertIn("| --- | --- |", parsed.text)
            self.assertIn(r"Mind\|Bridge", parsed.text)
            self.assertIn("第一行<br>第二行", parsed.text)
            ordered_fragments = [
                "# 中文标题",
                "标题后的普通段落。",
                "- 无序项目",
                "1. 有序项目",
                "| 名称 | 说明 |",
                "表格后的段落。",
            ]
            positions = [parsed.text.index(fragment) for fragment in ordered_fragments]
            self.assertEqual(positions, sorted(positions))
            self.assertNotIn("\n\n\n", parsed.text)

    def test_rejects_docx_over_uncompressed_limit(self):
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "large.docx"
            document = DocxDocument()
            document.add_paragraph("content")
            document.save(path)

            with self.assertRaisesRegex(DocumentParseError, "安全限制"):
                parse_document(
                    path,
                    path.name,
                    parser_settings(knowledge_docx_max_uncompressed_bytes=1),
                )

    def test_renders_a_zero_row_table_stably(self):
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "empty-table.docx"
            document = DocxDocument()
            document.add_table(rows=0, cols=2)
            document.save(path)

            parsed = parse_document(path, path.name, parser_settings())

            self.assertEqual(parsed.text, "表格（0 行，2 列）")


class PdfDocumentParserTests(unittest.TestCase):
    def test_preserves_pages_and_reports_empty_pages(self):
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "pages.pdf"
            write_text_pdf(
                path,
                [["Page one line one", "Page one line two"], None, ["Page three"]],
            )

            parsed = parse_document(path, path.name, parser_settings())

            self.assertEqual(parsed.parser_name, "pypdf")
            self.assertEqual(parsed.page_count, 3)
            self.assertEqual(parsed.metadata["empty_pages"], [2])
            self.assertIn("Page one line one\nPage one line two", parsed.text)
            self.assertIn("--- PDF PAGE 2 ---", parsed.text)
            self.assertIn("--- PDF PAGE 3 ---", parsed.text)
            self.assertIn("Page three", parsed.text)
            self.assertIn("PDF 第 2 页未提取到文本", parsed.warnings)

    def test_rejects_pdf_without_extractable_text(self):
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "scan.pdf"
            write_text_pdf(path, [None, None])

            with self.assertRaisesRegex(DocumentParseError, "扫描型 PDF.*不支持 OCR"):
                parse_document(path, path.name, parser_settings())


class RecursiveCharacterSplitterTests(unittest.TestCase):
    def test_validates_all_parameter_boundaries(self):
        self.assertEqual(validate_splitter_config(100, 0), SplitterConfig(100, 0))
        self.assertEqual(validate_splitter_config(4000, 1000), SplitterConfig(4000, 1000))

        invalid = [
            (99, 0, "recursive_character"),
            (4001, 0, "recursive_character"),
            (100, -1, "recursive_character"),
            (2000, 1001, "recursive_character"),
            (100, 100, "recursive_character"),
            (512, 64, "token"),
            (True, 0, "recursive_character"),
            (512, False, "recursive_character"),
        ]
        for chunk_size, overlap, splitter_type in invalid:
            with self.subTest(
                chunk_size=chunk_size,
                overlap=overlap,
                splitter_type=splitter_type,
            ), self.assertRaises(SplitterConfigError):
                validate_splitter_config(chunk_size, overlap, splitter_type)

    def test_splits_by_characters_and_preserves_markdown_and_chinese(self):
        text = (
            "# 第一章\n\n"
            + "这是第一段。它包含中文问号？也包含中文感叹号！" * 4
            + "\n\n## 列表\n\n- 第一项\n- 第二项\n\n"
            + "最后一段；继续说明。" * 5
        )
        config = validate_splitter_config(100, 20)

        chunks = split_text(text, config)

        self.assertGreater(len(chunks), 1)
        self.assertTrue(all(0 < len(chunk) <= 100 for chunk in chunks))
        self.assertTrue(any("# 第一章" in chunk for chunk in chunks))
        self.assertTrue(any("## 列表" in chunk for chunk in chunks))
        self.assertTrue(any("- 第一项" in chunk for chunk in chunks))
        self.assertTrue(any("。" in chunk and "？" in chunk for chunk in chunks))
        self.assertEqual(chunks, split_text(text, config))

    def test_empty_text_produces_no_chunks(self):
        self.assertEqual(split_text(" \n\t ", SplitterConfig(100, 0)), [])


if __name__ == "__main__":
    unittest.main()
