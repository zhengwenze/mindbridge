"""Focused unit tests for safe, streaming knowledge-document uploads."""

from __future__ import annotations

import tempfile
import unittest
from pathlib import Path
from types import SimpleNamespace

from docx import Document as DocxDocument

from app.services.knowledge import (
    KnowledgeDocumentInvalid,
    KnowledgeDocumentTooLarge,
    KnowledgeDocumentUnsupported,
    extract_document_chunks,
    normalize_relative_path,
    receive_upload,
    validate_document_extension,
)


def settings(root: Path, **overrides):
    values = {
        "project_root": root,
        "knowledge_upload_max_bytes": 50 * 1024 * 1024,
        "knowledge_upload_read_chunk_bytes": 1024 * 1024,
        "knowledge_docx_max_uncompressed_bytes": 200 * 1024 * 1024,
        "knowledge_chunk_size": 32,
        "knowledge_chunk_overlap": 4,
    }
    values.update(overrides)
    return SimpleNamespace(**values)


class FakeUpload:
    def __init__(self, filename: str, data: bytes):
        self.filename = filename
        self.data = data
        self.offset = 0
        self.read_sizes: list[int] = []

    async def read(self, size: int) -> bytes:
        self.read_sizes.append(size)
        block = self.data[self.offset:self.offset + size]
        self.offset += len(block)
        return block


def write_text_pdf(path: Path, text: str) -> None:
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"BT /F1 12 Tf 72 100 Td ({escaped}) Tj ET".encode("ascii")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 200] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        f"<< /Length {len(stream)} >>\nstream\n".encode("ascii") + stream + b"\nendstream",
    ]
    output = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, body in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{index} 0 obj\n".encode("ascii"))
        output.extend(body)
        output.extend(b"\nendobj\n")
    xref_offset = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode("ascii")
    )
    path.write_bytes(output)


class DocumentPathTests(unittest.TestCase):
    def test_normalizes_folder_paths_and_rejects_traversal(self):
        self.assertEqual(normalize_relative_path("制度/学生 手册.pdf", "ignored.pdf"), "制度/学生 手册.pdf")
        self.assertEqual(normalize_relative_path(None, "notes.md"), "notes.md")
        for invalid in ("../secret.txt", "/tmp/a.txt", "C:\\temp\\a.txt", "folder//a.txt", "folder/./a.txt"):
            with self.subTest(invalid=invalid), self.assertRaises(KnowledgeDocumentInvalid):
                normalize_relative_path(invalid, "fallback.txt")

    def test_rejects_unsupported_extension(self):
        validate_document_extension("guide.docx")
        with self.assertRaises(KnowledgeDocumentUnsupported):
            validate_document_extension("legacy.doc")


class StreamingUploadTests(unittest.IsolatedAsyncioTestCase):
    async def test_reads_only_configured_blocks(self):
        with tempfile.TemporaryDirectory() as directory:
            upload = FakeUpload("notes.txt", b"0123456789")
            config = settings(
                Path(directory),
                knowledge_upload_read_chunk_bytes=4,
                knowledge_upload_max_bytes=10,
            )
            path, size = await receive_upload(upload, config)
            self.assertEqual(size, 10)
            self.assertEqual(path.read_bytes(), b"0123456789")
            self.assertTrue(all(value == 4 for value in upload.read_sizes))
            path.unlink()

    async def test_removes_partial_file_when_limit_is_exceeded(self):
        with tempfile.TemporaryDirectory() as directory:
            upload = FakeUpload("large.txt", b"123456")
            config = settings(
                Path(directory),
                knowledge_upload_read_chunk_bytes=4,
                knowledge_upload_max_bytes=5,
            )
            with self.assertRaises(KnowledgeDocumentTooLarge):
                await receive_upload(upload, config)
            temp_root = Path(directory) / "data" / "knowledge-files" / ".tmp"
            self.assertEqual(list(temp_root.iterdir()), [])


class DocumentExtractionTests(unittest.TestCase):
    def test_extracts_all_supported_formats(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            config = settings(root, knowledge_chunk_size=512, knowledge_chunk_overlap=64)
            samples = {
                "notes.txt": "plain text upload",
                "guide.md": "markdown upload guide",
                "readme.markdown": "long markdown upload",
            }
            for name, content in samples.items():
                path = root / name
                path.write_text(content, encoding="utf-8")
                self.assertIn(content, " ".join(extract_document_chunks(path, name, config)))

            pdf_path = root / "guide.pdf"
            write_text_pdf(pdf_path, "PDF upload guide")
            self.assertIn("PDF upload guide", " ".join(extract_document_chunks(pdf_path, pdf_path.name, config)))

            docx_path = root / "guide.docx"
            document = DocxDocument()
            document.add_paragraph("DOCX upload guide")
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "folder"
            table.cell(0, 1).text = "document"
            document.save(docx_path)
            extracted = " ".join(extract_document_chunks(docx_path, docx_path.name, config))
            self.assertIn("DOCX upload guide", extracted)
            self.assertIn("| folder | document |", extracted)

    def test_rejects_invalid_or_empty_documents(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            config = settings(root)
            invalid_text = root / "invalid.txt"
            invalid_text.write_bytes(b"\xff\xfe\xfa")
            with self.assertRaises(KnowledgeDocumentInvalid):
                list(extract_document_chunks(invalid_text, invalid_text.name, config))

            invalid_pdf = root / "invalid.pdf"
            invalid_pdf.write_bytes(b"not a pdf")
            with self.assertRaises(KnowledgeDocumentInvalid):
                list(extract_document_chunks(invalid_pdf, invalid_pdf.name, config))

            empty = root / "empty.md"
            empty.write_text("   \n", encoding="utf-8")
            self.assertEqual(list(extract_document_chunks(empty, empty.name, config)), [])


if __name__ == "__main__":
    unittest.main()
